"""
doc_generator.py — Generate .docx documents for tender preparation using GPT-4o + python-docx.
Only generates declarations, letters, and undertakings — never government-issued documents.
"""

from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI

GENERATION_PROMPT = """You are a professional document writer for Indian Government procurement (GeM — Government e-Marketplace).

Generate the body text ONLY for a formal "{document_type}" document.

Company Details:
{company_details}

Tender Details:
{tender_details}

Rules:
- Write 2–4 short formal paragraphs
- Include company name, tender number, and department where relevant
- Do NOT write a title, subject line, salutation, "To:", "From:", or signature block — those are added separately
- Do NOT use bullet points
- Return ONLY the body paragraph text, nothing else
"""


def _fmt_company(p: dict) -> str:
    fields = [
        ("Company Name",         p.get("company_name")),
        ("Address",              p.get("address")),
        ("GST Number",           p.get("gst_number")),
        ("PAN Number",           p.get("pan_number")),
        ("MSME/Udyam Number",    p.get("msme_number")),
        ("Bank Name",            p.get("bank_name")),
        ("Account Number",       p.get("account_number")),
        ("IFSC Code",            p.get("ifsc_code")),
        ("Authorized Signatory", p.get("authorized_signatory_name")),
        ("Designation",          p.get("authorized_signatory_designation")),
    ]
    return "\n".join(f"{k}: {v}" for k, v in fields if v)


def _fmt_tender(t: dict) -> str:
    fields = [
        ("Tender Number",  t.get("tender_number")),
        ("Date",           t.get("date")),
        ("Department",     t.get("department_name")),
        ("Organisation",   t.get("organization_name")),
        ("Bid End Date",   t.get("bid_end_datetime")),
    ]
    return "\n".join(f"{k}: {v}" for k, v in fields if v)


def _is_image(path: str) -> bool:
    return Path(path).suffix.lower() in {".png", ".jpg", ".jpeg"}


def _add_image_or_space(doc, path: str | None, width_inches: float = 1.5, blank_lines: int = 2):
    """Add an image paragraph if the file exists and is a valid image, else blank lines."""
    if path and Path(path).exists() and _is_image(path):
        try:
            para = doc.add_paragraph()
            para.add_run().add_picture(path, width=Inches(width_inches))
            return
        except Exception:
            pass
    for _ in range(blank_lines):
        doc.add_paragraph("")


def generate_ai_document(document_type: str, profile: dict, tender: dict, output_path: str) -> str:
    """Call GPT-4o to generate document text and save as a formatted .docx file."""
    client = OpenAI()
    prompt = GENERATION_PROMPT.format(
        document_type=document_type,
        company_details=_fmt_company(profile),
        tender_details=_fmt_tender(tender),
    )
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1024,
    )
    body_text = response.choices[0].message.content.strip()

    doc = Document()

    # Title
    title = doc.add_paragraph()
    run = title.add_run(document_type.upper())
    run.bold = True
    run.font.size = Pt(13)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Tender reference block
    doc.add_paragraph(f"Tender No : {tender.get('tender_number', '')}")
    doc.add_paragraph(f"Department: {tender.get('department_name', '')}")
    doc.add_paragraph("")
    doc.add_paragraph("To,")
    doc.add_paragraph(f"{tender.get('organization_name', '')}")
    doc.add_paragraph("")

    # Generated body
    for para in body_text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # Signature block
    doc.add_paragraph("")
    doc.add_paragraph("Yours faithfully,")
    doc.add_paragraph("")
    doc.add_paragraph(f"For {profile.get('company_name', '')}")
    doc.add_paragraph("")
    _add_image_or_space(doc, profile.get("signature_file_path"), width_inches=1.5)
    doc.add_paragraph("Authorized Signatory")
    doc.add_paragraph(profile.get("authorized_signatory_name", ""))
    doc.add_paragraph(profile.get("authorized_signatory_designation", ""))
    doc.add_paragraph(f"GST No : {profile.get('gst_number', '')}")
    doc.add_paragraph("")
    _add_image_or_space(doc, profile.get("stamp_file_path"), width_inches=1.5)
    doc.add_paragraph(f"Date : {datetime.now().strftime('%d-%m-%Y')}")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
