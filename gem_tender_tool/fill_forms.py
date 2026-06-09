"""
fill_forms.py — fill {{placeholder}} tokens in a .docx template.

Word often splits a single {{token}} across multiple runs, so naive
run-by-run replacement misses them. We replace at the paragraph level by
flattening the paragraph text, substituting, then rewriting it into a single run
(formatting of the first run is preserved, which is fine for declaration forms).
"""

import re
import datetime
from docx import Document

TOKEN = re.compile(r"\{\{(\w+)\}\}")


def _fill_paragraph(paragraph, context):
    if "{{" not in paragraph.text:
        return
    new_text = TOKEN.sub(
        lambda m: str(context.get(m.group(1), m.group(0))), paragraph.text)
    if new_text == paragraph.text:
        return
    # keep first run's formatting, drop the rest, write flattened text
    for run in paragraph.runs[1:]:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def fill_template(template_path, output_path, context):
    doc = Document(template_path)
    for p in doc.paragraphs:
        _fill_paragraph(p, context)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _fill_paragraph(p, context)
    doc.save(output_path)
    return output_path


def build_context(company_profile, tender):
    """Merge company profile + selected tender fields into one flat dict."""
    ctx = dict(company_profile)
    ctx.update({
        "bid_number": tender.get("bid_number", ""),
        "bid_dated": tender.get("bid_dated", ""),
        "boq_title": tender.get("boq_title", ""),
        "organisation": tender.get("organisation", ""),
        "today": datetime.date.today().strftime("%d-%m-%Y"),
    })
    return ctx
